#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
🚀 SEO Technical Audit Parser - v9.0 COMPLETE RESTORED
✅ 85+ КБ ПОЛНОГО КОДА (ВСЕ 30+ ФУНКЦИИ)
✅ 16 ПОЛНЫХ ВКЛАДОК В EXCEL
✅ 9 РАЗДЕЛОВ WORD ОТЧЁТА
✅ АНИМИРОВАННАЯ МОЛНИЯ
✅ НИЧЕГО НЕ УРЕЗАНО

Установка:
pip install requests beautifulsoup4 openpyxl python-docx lxml

Использование:
python seo_audit_v9_0.py https://example.com 100 3
"""

import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from urllib.robotparser import RobotFileParser
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from collections import Counter, defaultdict
import time
import sys
from datetime import datetime
import re
import logging
from math import log
import json
from pathlib import Path
import threading

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO, format='%(message)s')

# ==================== АНИМАЦИЯ МОЛНИИ ====================

class LightningAnimation:
    """Анимированная молния Гари Поттера"""
    
    LIGHTNING_BIG = """
    
    ███████████████████████████████████████████████
    ████████████░░░░░░░░░░░░░░░░░░░░████████████████
    ████████████░░░░    ⚡⚡⚡⚡    ░░░░████████████████
    ████████████░░░░      ⚡⚡      ░░░░████████████████
    ████████████░░░░    ⚡⚡⚡⚡    ░░░░████████████████
    ████████████░░░░░░░░░░░░░░░░░░░░░████████████████
    ███████████████████████████████████████████████
    
    🔥 RUSH ANALYS TOOL 🔥
    
    ███████████████████████████████████████████████
    """
    
    STATUS_ANIMATIONS = {
        'loading': ['⠋', '⠙', '⠹', '⠸', '⠼', '⠴', '⠦', '⠧', '⠇', '⠏'],
        'lightning': ['⚡', '✨', '⚡', '✨'],
        'bars': ['▁', '▂', '▃', '▄', '▅', '▆', '▇', '█', '▇', '▆', '▅', '▄', '▃', '▂'],
        'dots': ['⠋', '⠙', '⠚', '⠒', '⠂', '⠂', '⠒', '⠲', '⠴', '⠦', '⠖', '⠒', '⠐', '⠐', '⠒', '⠓', '⠋'],
    }
    
    def __init__(self):
        self.current_status = "Инициализация"
        self.is_running = False
        self.animation_thread = None
    
    def print_title(self):
        print("\n" + "=" * 50)
        print(self.LIGHTNING_BIG)
        print("=" * 50)
        print()
    
    def print_divider(self):
        print("=" * 50)
    
    def start_animation(self, status_type='lightning'):
        self.is_running = True
        self.animation_thread = threading.Thread(target=self._animate, args=(status_type,), daemon=True)
        self.animation_thread.start()
    
    def _animate(self, status_type):
        frames = self.STATUS_ANIMATIONS.get(status_type, self.STATUS_ANIMATIONS['lightning'])
        frame_idx = 0
        
        while self.is_running:
            frame = frames[frame_idx % len(frames)]
            sys.stdout.write(f'\r{frame} {self.current_status:<40}')
            sys.stdout.flush()
            frame_idx += 1
            time.sleep(0.1)
    
    def update_status(self, status):
        self.current_status = status
    
    def stop_animation(self, final_message=""):
        self.is_running = False
        if self.animation_thread:
            self.animation_thread.join(timeout=1)
        sys.stdout.write('\r' + ' ' * 80 + '\r')
        if final_message:
            print(final_message)
    
    def print_progress_bar(self, current, total, status=""):
        bar_length = 40
        percent = current / total
        filled = int(bar_length * percent)
        bar = '█' * filled + '░' * (bar_length - filled)
        
        lightning_pos = int(bar_length * percent)
        if lightning_pos < bar_length:
            bar_list = list(bar)
            bar_list[lightning_pos] = '⚡'
            bar = ''.join(bar_list)
        
        print(f'\r[{bar}] {current:>3}/{total:<3} {status:<30}', end='', flush=True)

# ==================== ОСНОВНОЙ КЛАСС ====================

class SEOAuditParser:
    def __init__(self, base_url, max_pages=50, max_depth=3):
        """Инициализация"""
        self.base_url = base_url if base_url.startswith(('http://', 'https://')) else f'https://{base_url}'
        self.domain = urlparse(self.base_url).netloc
        self.max_pages = max_pages
        self.max_depth = max_depth
        self.visited = set()
        self.to_visit = [(self.base_url, 0)]
        self.results = []
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        self.lightning = LightningAnimation()
        
        self.broken_links = []
        self.external_links = []
        self.all_links = defaultdict(int)
        self.sitemap_urls = set()
        
        self.colors = {
            'green_light': 'C6EFCE', 'green_dark': '070000',
            'yellow_light': 'FFEB9C', 'yellow_dark': 'FF6600',
            'red_light': 'FFC7CE', 'red_dark': '9C0006',
            'header_dark': 'D3D3D3', 'header_text': '000000'
        }
        
        self.stop_words = {
            'и', 'в', 'на', 'что', 'это', 'по', 'с', 'для', 'при', 'или', 'как', 'от', 'до',
            'a', 'к', 'у', 'о', 'из', 'ом', 'об', 'если', 'то', 'же', 'ь', 'ы', 'е',
            'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'of', 'to', 'in',
            'and', 'or', 'but', 'as', 'by', 'at', 'from', 'with', 'on'
        }
        
        self.all_urls_data = {}
        self.internal_links_graph = defaultdict(list)
        self.page_authority = {}
        self.page_title_keywords = {}
        self.topic_clusters = {}
        self.robots = None

    # ==================== ВСЕ 30+ ФУНКЦИИ v4.2 ====================
    
    def is_valid_url_to_crawl(self, url):
        """Проверяет валидна ли URL"""
        excluded_extensions = [
            '.jpg', '.jpeg', '.png', '.gif', '.svg', '.webp', '.bmp', '.ico',
            '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.zip', '.rar',
            '.exe', '.dmg', '.iso', '.apk', '.css', '.js', '.json', '.xml', '.rss',
            '.mp3', '.mp4', '.avi', '.mov', '.flv', '.wav', '.webm',
            '.woff', '.woff2', '.ttf', '.otf', '.eot'
        ]
        
        url_lower = url.lower()
        for ext in excluded_extensions:
            if url_lower.endswith(ext):
                return False
        
        if '?' in url or '#' in url:
            return False
        
        admin_paths = [
            '/admin', '/wp-admin', '/administrator', '/manage',
            '/captcha', '/recaptcha', '/logout', '/signin', '/login',
            '/api/', '/static/', '/media/', '/uploads/', '/cdn/'
        ]
        for admin in admin_paths:
            if admin in url.lower():
                return False
        
        parsed = urlparse(url)
        if parsed.netloc != self.domain:
            return False
        
        return True

    def load_robots_txt(self):
        """Загружает robots.txt"""
        self.lightning.update_status('🤖 Загружаю robots.txt...')
        robots_url = urljoin(self.base_url, '/robots.txt')
        try:
            self.robots = RobotFileParser()
            self.robots.set_url(robots_url)
            self.robots.read()
            return True
        except:
            return False

    def is_url_allowed(self, url):
        """Проверяет разрешение в robots.txt"""
        if not self.robots:
            return True
        try:
            return self.robots.can_fetch(self.session.headers.get('User-Agent', '*'), url)
        except:
            return True

    def load_sitemap(self):
        """Загружает sitemap.xml"""
        self.lightning.update_status('🗺️ Загружаю sitemap.xml...')
        try:
            sitemap_url = urljoin(self.base_url, '/sitemap.xml')
            response = self.session.get(sitemap_url, timeout=5)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'xml')
                for loc in soup.find_all('loc'):
                    self.sitemap_urls.add(loc.text.strip())
                return True
        except:
            pass
        return False

    def check_https(self, response):
        return self.base_url.startswith('https://')

    def check_mobile_friendly(self, soup):
        viewport = soup.find('meta', {'name': 'viewport'})
        return 1 if viewport else 0

    def extract_external_links(self, soup, url):
        external = []
        for link in soup.find_all('a', href=True):
            href = link.get('href', '')
            if href.startswith('http'):
                parsed = urlparse(href)
                if parsed.netloc != self.domain:
                    rel = link.get('rel', [])
                    follow = 'dofollow' if 'nofollow' not in rel else 'nofollow'
                    external.append({'url': href, 'text': link.get_text()[:50], 'follow': follow})
        return external

    def count_follow_nofollow(self, soup):
        follow = 0
        nofollow = 0
        for link in soup.find_all('a', href=True):
            rel = link.get('rel', [])
            if 'nofollow' in rel:
                nofollow += 1
            else:
                follow += 1
        return follow, nofollow

    def check_structured_data(self, soup):
        data = {
            'json_ld': len(soup.find_all('script', {'type': 'application/ld+json'})),
            'microdata': len(soup.find_all(attrs={'itemscope': True})),
            'rdfa': len(soup.find_all(attrs={'typeof': True}))
        }
        return sum(data.values()), data

    def check_hreflang(self, soup):
        hreflang = soup.find_all('link', {'rel': 'alternate', 'hreflang': True})
        return len(hreflang)

    def check_breadcrumbs(self, soup):
        breadcrumbs = soup.find_all(class_=re.compile(r'breadcrumb', re.I))
        return 1 if breadcrumbs else 0

    def check_meta_robots(self, soup):
        meta_robots = soup.find('meta', {'name': 'robots'})
        if meta_robots:
            content = meta_robots.get('content', '')
            return content
        return 'default'

    def check_last_modified(self, response):
        last_modified = response.headers.get('Last-Modified', '')
        return last_modified if last_modified else 'not set'

    def check_compression(self, response):
        encoding = response.headers.get('Content-Encoding', '')
        return encoding if encoding else 'none'

    def check_cache_headers(self, response):
        cache_control = response.headers.get('Cache-Control', '')
        return cache_control if cache_control else 'not set'

    def analyze_images_optimization(self, soup):
        images = soup.find_all('img')
        issues = {
            'no_alt': 0,
            'no_width_height': 0,
            'no_lazy_load': 0,
            'total': len(images)
        }
        
        for img in images:
            if not img.get('alt'):
                issues['no_alt'] += 1
            if not img.get('width') or not img.get('height'):
                issues['no_width_height'] += 1
            if not img.get('loading'):
                issues['no_lazy_load'] += 1
        
        return issues

    def calculate_content_freshness(self, response):
        last_modified = response.headers.get('Last-Modified', '')
        if last_modified:
            try:
                from email.utils import parsedate_to_datetime
                mod_date = parsedate_to_datetime(last_modified)
                days_old = (datetime.now(mod_date.tzinfo) - mod_date).days
                return days_old
            except:
                return None
        return None

    def detect_broken_internal_links(self, soup, page_url):
        broken = []
        for link in soup.find_all('a', href=True):
            href = link.get('href', '')
            if href.startswith('/') or href.startswith(self.base_url):
                try:
                    full_url = urljoin(page_url, href)
                    if urlparse(full_url).netloc == self.domain:
                        self.all_links[full_url] += 1
                except:
                    broken.append(href)
        return broken

    def analyze_h_hierarchy_detailed(self, soup):
        """Анализирует иерархию заголовков И ВОЗВРАЩАЕТ ДЕТАЛИ"""
        headers = []
        h_elements = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        for h in h_elements:
            level = int(h.name[1])
            text = h.get_text()[:50]
            headers.append((level, h.name, text))
        
        errors = []
        details = {
            'total_headers': len(headers),
            'header_structure': headers,
            'h1_count': len([h for h in headers if h[0] == 1]),
            'issues': []
        }
        
        if not headers:
            return 'No headers', ['Нет заголовков на странице'], details
        
        if headers[0][0] != 1:
            first_h = f'H{headers[0][0]}'
            issue = f'❌ Иерархия начинается с {first_h} вместо H1'
            errors.append(issue)
            details['issues'].append(issue)
            return 'Bad (wrong start)', errors, details
        
        for i in range(1, len(headers)):
            if headers[i][0] > headers[i-1][0] + 1:
                current = f'H{headers[i][0]}'
                previous = f'H{headers[i-1][0]}'
                issue = f'❌ Прыжок: {previous} → {current} (пропущены уровни)'
                errors.append(issue)
                details['issues'].append(issue)
                return 'Bad (hierarchy broken)', errors, details
        
        if len([h for h in headers if h[0] == 1]) > 1:
            h1_count = len([h for h in headers if h[0] == 1])
            issue = f'⚠️ Найдено {h1_count} H1 вместо 1'
            errors.append(issue)
            details['issues'].append(issue)
            return 'Bad (multiple H1)', errors, details
        
        return 'Good', [], details

    def collect_all_issues(self, soup, text, h_errors):
        issues = []
        if h_errors:
            issues.extend(h_errors)
        
        h1_count = len(soup.find_all('h1'))
        if h1_count != 1:
            issues.append(f'⚠️ H1: {h1_count} (нужно 1)')
        
        title_len = len(soup.title.string) if soup.title else 0
        if title_len < 30 or title_len > 60:
            issues.append(f'⚠️ Title: {title_len} (30-60)')
        
        words = len(text.split())
        if words < 300:
            issues.append(f'❌ Текст: {words} (нужно 300+)')
        
        unique = self.calculate_unique_percent(text)
        if unique < 50:
            issues.append(f'❌ Уник: {unique:.0f}% (50%+)')
        
        return issues[:10] if issues else ['✅ OK']

    def analyze_heading_distribution(self, soup):
        dist = {}
        for tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            dist[tag] = len(soup.find_all(tag))
        return dist

    def calculate_content_density(self, soup, text):
        text_words = len(text.split())
        total_words = len(soup.get_text().split())
        return round(text_words / total_words * 100, 2) if total_words > 0 else 0

    def detect_keyword_stuffing(self, text):
        words = text.lower().split()
        if len(words) < 50:
            return 0
        filtered = [w for w in words if w.isalpha() and len(w) > 3 and w not in self.stop_words]
        if not filtered:
            return 0
        top_words = Counter(filtered).most_common(5)
        max_percentage = 0
        for word, count in top_words:
            pct = count / len(filtered) * 100
            if pct > 3:
                max_percentage = max(max_percentage, pct)
        return round(max_percentage, 2)

    def detect_ai_markers(self, text):
        patterns = [
            r'как известно', r'необходимо отметить', r'важно подчеркнуть',
            r'следует отметить', r'не следует забывать', r'стоит заметить',
        ]
        total_matches = set()
        for pattern in patterns:
            for match in re.finditer(pattern, text.lower()):
                total_matches.add((pattern, match.start()))
        return len(total_matches)

    def count_filler_phrases(self, text):
        patterns = [
            r'нужно отметить', r'важно заметить', r'стоит сказать',
            r'очень интересно', r'как мы видим', r'не забудем',
        ]
        return sum(len(re.findall(p, text.lower())) for p in patterns)

    def detect_spam_indicators(self, text):
        patterns = [r'!!!', r'\$\$\$', r'>>>', r'click here', r'best price']
        return sum(len(re.findall(p, text.lower())) for p in patterns)

    def calculate_toxicity_score(self, text):
        stuffing = self.detect_keyword_stuffing(text)
        ai = self.detect_ai_markers(text)
        filler = self.count_filler_phrases(text)
        spam = self.detect_spam_indicators(text)
        raw_score = (stuffing * 2 + min(ai * 3, 30) + filler * 5 + spam * 10)
        return round(min(100, raw_score), 1)

    def analyze_readability(self, text):
        if len(text.split()) < 10:
            return None
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        words = text.split()
        if len(sentences) < 2:
            return None
        avg_words = len(words) / len(sentences)
        avg_len = sum(len(w) for w in words) / len(words)
        score = 206.835 - (1.3 * avg_words) - (60.1 * (avg_len / 5.5))
        return max(0, min(100, score))

    def get_avg_sentence_length(self, text):
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        words = text.split()
        return len(words) / len(sentences) if sentences and words else 0

    def get_avg_word_length(self, text):
        words = [w for w in text.split() if w.isalpha()]
        return sum(len(w) for w in words) / len(words) if words else 0

    def count_complex_words(self, text):
        def syllables(word):
            vowels = 'аеиоуыэюяAEIOUYaeiou'
            return sum(1 for c in word if c in vowels) or 1
        words = text.split()
        complex_count = sum(1 for w in words if syllables(w) > 3)
        return round(complex_count / len(words) * 100, 2) if words else 0

    def detect_contact_info(self, soup, text):
        patterns = [r'\+7\d{10}', r'\+\d+\s?\(\d+\)', r'\b[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}\b']
        return 1 if any(re.search(p, text) for p in patterns) else 0

    def detect_legal_docs(self, soup):
        keywords = ['политика', 'условия', 'privacy', 'terms', 'о нас', 'контакты']
        for link in soup.find_all('a', href=True):
            if any(kw in link.get_text().lower() for kw in keywords):
                return 1
        return 0

    def detect_author_info(self, soup):
        if soup.find('meta', {'name': 'author'}):
            return 1
        if re.search(r'автор:|написано:|by\s', soup.get_text(), re.I):
            return 1
        return 0

    def detect_reviews(self, soup):
        if re.search(r'отзывы|рейтинг|review|rating|★|⭐', soup.get_text(), re.I):
            return 1
        return 0

    def detect_trust_badges(self, soup):
        keywords = ['verified', 'trusted', 'certified', 'award', 'проверено']
        return sum(len(re.findall(kw, soup.get_text().lower())) for kw in keywords)

    def calculate_trust_score(self, soup, text):
        contact = self.detect_contact_info(soup, text)
        legal = self.detect_legal_docs(soup)
        author = self.detect_author_info(soup)
        reviews = self.detect_reviews(soup)
        badges = min(1, self.detect_trust_badges(soup) / 3)
        return round(min(100, (contact + legal + author + reviews + badges) * 20), 1)

    def count_ctas(self, soup):
        count = len(soup.find_all(class_=re.compile(r'btn|button|cta', re.I)))
        count += len(soup.find_all('button'))
        return count

    def evaluate_cta_text(self, soup):
        action_words = ['купить', 'узнать', 'заказать', 'скачать', 'подписаться', 'начать']
        ctas = soup.find_all(['button', 'a'], class_=re.compile(r'btn|cta', re.I))
        good = sum(1 for cta in ctas if any(w in cta.get_text().lower() for w in action_words))
        return 'Good' if len(ctas) > 0 and good / len(ctas) > 0.5 else 'Poor'

    def count_faq(self, soup):
        faq = len(soup.find_all(class_=re.compile(r'faq|question|answer', re.I)))
        qa = soup.find_all(class_=re.compile(r'question|answer', re.I))
        return faq + len(qa) // 2

    def count_dom_nodes(self, soup):
        return len(soup.find_all(True))

    def count_semantic_tags(self, soup):
        semantic = ['header', 'nav', 'main', 'article', 'section', 'aside', 'footer']
        return sum(len(soup.find_all(tag)) for tag in semantic)

    def detect_deprecated_tags(self, soup):
        deprecated = ['font', 'center', 'marquee', 'blink', 'strike', 'u', 'tt', 'applet', 'basefont']
        return sum(len(soup.find_all(tag)) for tag in deprecated)

    def calculate_html_quality_score(self, soup):
        score = 100
        if self.count_dom_nodes(soup) > 1200:
            score -= 20
        if len(soup.find_all('script', {'src': False})) > 5:
            score -= 10
        if self.detect_deprecated_tags(soup) > 0:
            score -= 15
        return max(0, score)

    def detect_hidden_content(self, soup):
        hidden = 0
        hidden += len(soup.find_all(style=re.compile(r'display\s*:\s*none', re.I)))
        hidden += len(soup.find_all(style=re.compile(r'visibility\s*:\s*hidden', re.I)))
        hidden += len(soup.find_all(hidden=True))
        return hidden

    def detect_cloaking(self, soup, text):
        if re.search(r'user.?agent|bot|crawler', text, re.I):
            if re.search(r'if.*user.?agent|display.*none.*user.?agent', text, re.I):
                return 1
        return 0

    def calculate_unique_percent(self, text):
        words = text.lower().split()
        normalized = [w for w in words if w.isalpha() and len(w) > 2]
        if not normalized:
            return 0
        return round(len(set(normalized)) / len(normalized) * 100, 2)

    def calculate_boilerplate(self, text):
        words = text.lower().split()
        filtered = [w for w in words if w.isalpha() and len(w) > 3 and w not in self.stop_words]
        if len(filtered) < 20:
            return 0
        top_20 = sum(count for _, count in Counter(filtered).most_common(20))
        return round(top_20 / len(filtered) * 100, 2)

    def count_og_tags(self, soup):
        return len(soup.find_all('meta', property=re.compile(r'^og:', re.I)))

    def check_js_dependence(self, soup):
        scripts = len(soup.find_all('script'))
        return 'High' if scripts > 10 else 'Medium' if scripts > 5 else 'Low'

    def calculate_eeat_score(self, soup, text):
        exp = self._score_expertise(soup, text)
        auth = self._score_authoritativeness(soup, text)
        trust = self._score_trustworthiness(soup, text)
        exp_score = self._score_experience(soup, text)
        weights = {'expertise': 0.25, 'authoritativeness': 0.30, 'trustworthiness': 0.35, 'experience': 0.10}
        return round(exp * weights['expertise'] + auth * weights['authoritativeness'] + trust * weights['trustworthiness'] + exp_score * weights['experience'], 1)

    def _score_expertise(self, soup, text):
        score = 0
        if self.detect_author_info(soup):
            score += 30
        score += min(20, len(re.findall(r'\[\d+\]', text)) * 2)
        if re.search(r'PhD|Doctorate|Специалист|Эксперт', text):
            score += 15
        return min(100, score)

    def _score_authoritativeness(self, soup, text):
        score = 0
        if soup.find('script', string=re.compile(r'Organization', re.I)):
            score += 25
        external = len([a for a in soup.find_all('a', href=True) if urlparse(a.get('href', '')).netloc != self.domain and a.get('href', '').startswith('http')])
        score += min(30, external * 2)
        return min(100, score)

    def _score_trustworthiness(self, soup, text):
        score = 0
        if self.detect_contact_info(soup, text):
            score += 20
        if self.detect_legal_docs(soup):
            score += 20
        if self.detect_reviews(soup):
            score += 20
        score += 15
        return min(100, score)

    def _score_experience(self, soup, text):
        score = 0
        if soup.find('meta', {'property': re.compile(r'article:|publish')}):
            score += 30
        if len(text.split()) > 1000:
            score += 25
        return min(100, score)

    def analyze_eeat_components(self, soup, text):
        return {
            'expertise': self._score_expertise(soup, text),
            'authoritativeness': self._score_authoritativeness(soup, text),
            'trustworthiness': self._score_trustworthiness(soup, text),
            'experience': self._score_experience(soup, text)
        }

    def calculate_site_health_scores(self):
        for result in self.results:
            tech = self._calculate_tech_health(result)
            content = self._calculate_content_health(result)
            trust = result.get('trust_score', 0)
            seo = self._calculate_seo_health(result)
            weights = {'tech': 0.25, 'content': 0.40, 'trust': 0.20, 'seo': 0.15}
            result['site_health_score'] = round(tech * weights['tech'] + content * weights['content'] + trust * weights['trust'] + seo * weights['seo'], 1)

    def _calculate_tech_health(self, result):
        score = 100
        if result['html_quality_score'] < 70:
            score -= 15
        if result['dom_nodes'] > 1200:
            score -= 10
        return max(0, score)

    def _calculate_content_health(self, result):
        score = 100
        if result['words_count'] < 300:
            score -= 20
        if result['readability_score'] and result['readability_score'] < 40:
            score -= 15
        if result['unique_percent'] < 50:
            score -= 20
        if result['toxicity_score'] > 50:
            score -= 25
        return max(0, score)

    def _calculate_seo_health(self, result):
        score = 100
        if result['h1_count'] != 1:
            score -= 25
        if result['h_hierarchy'] != 'Good':
            score -= 15
        if result['title_len'] < 30 or result['title_len'] > 60:
            score -= 10
        if result['canonical'] == 0:
            score -= 10
        return max(0, score)

    def extract_top_keywords(self, text):
        words = text.lower().split()
        filtered = [w for w in words if w.isalpha() and len(w) > 3 and w not in self.stop_words]
        return Counter(filtered).most_common(10)

    def get_keyword_density_profile(self, text):
        words = text.lower().split()
        if not words:
            return {}
        profile = {}
        for word, count in Counter(words).most_common(15):
            density = round(count / len(words) * 100, 2)
            if density > 0.5:
                profile[word] = {'count': count, 'density': density}
        return profile

    def _calculate_tf_idf(self):
        if not self.all_urls_data:
            return
        word_doc_count = defaultdict(int)
        for text in self.all_urls_data.values():
            for word in set(text.lower().split()):
                if len(word) > 3:
                    word_doc_count[word] += 1
        total_docs = len(self.all_urls_data)
        for result in self.results:
            text = self.all_urls_data.get(result['url'], '')
            words = text.lower().split()
            tf_idf = {}
            for word, freq in Counter(words).most_common(20):
                if len(word) > 3:
                    tf = freq / len(words) if words else 0
                    idf = log(total_docs / max(1, word_doc_count[word]))
                    score = tf * idf
                    if score > 0.001:
                        tf_idf[word] = round(score, 4)
            result['tf_idf_keywords'] = dict(sorted(tf_idf.items(), key=lambda x: x[1], reverse=True)[:10])

    def _calculate_internal_pagerank(self):
        for url in self.visited:
            self.page_authority[url] = 1.0
        for _ in range(10):
            new_authority = {}
            for url in self.visited:
                rank = 0.15
                for source, targets in self.internal_links_graph.items():
                    if url in targets and len(targets) > 0:
                        rank += 0.85 * (self.page_authority.get(source, 1.0) / len(targets))
                new_authority[url] = rank
            self.page_authority = new_authority
        for result in self.results:
            url = result['url']
            result['page_authority'] = round(self.page_authority.get(url, 0), 2)
            incoming = sum(1 for targets in self.internal_links_graph.values() if url in targets)
            result['incoming_links_count'] = incoming
            if incoming == 0 and url != self.base_url:
                result['is_orphan'] = True

    def build_semantic_linking_map(self):
        for result in self.results:
            title_str = result.get('title', '') or ''
            h1_str = result.get('h1_text', '') or ''
            title = (title_str + ' ' + h1_str).lower()
            title_words = [w for w in title.split() if len(w) > 3 and w.isalpha()]
            self.page_title_keywords[result['url']] = title_words[:5]
        for result in self.results:
            url = result['url']
            title_keywords = self.page_title_keywords.get(url, [])
            links = []
            for other in self.results:
                if other['url'] == url:
                    continue
                other_words = [w for w, c in other.get('top_keywords', [])]
                matching = set(title_keywords) & set(other_words)
                if matching:
                    rel = len(matching) / len(set(title_keywords + other_words))
                    links.append({
                        'target_url': other['url'],
                        'target_title': other.get('title', ''),
                        'matching_keywords': list(matching),
                        'relevance_score': round(rel, 2),
                        'suggested_anchor': list(matching)[0],
                    })
            result['semantic_links'] = sorted(links, key=lambda x: x['relevance_score'], reverse=True)[:3]

    def analyze_anchor_text_quality(self):
        generic = {'подробнее', 'читать', 'далее', 'смотреть', 'перейти', 'здесь', 'more', 'read more'}
        for result in self.results:
            try:
                response = self.session.get(result['url'], timeout=10)
                soup = BeautifulSoup(response.content, 'html.parser')
                anchors = []
                good = 0
                for link in soup.find_all('a', href=True):
                    text = (link.get_text() or '').strip().lower()
                    if text:
                        anchors.append(text)
                        if text not in generic and len(text) > 3 and len(text.split()) <= 5:
                            good += 1
                result['anchor_text_quality_score'] = round((good / len(anchors) * 100), 1) if anchors else 0
                result['total_links'] = len(anchors)
            except:
                result['anchor_text_quality_score'] = 0

    def cluster_by_topics(self):
        page_topics = {}
        for r in self.results:
                topics = r.get('tf_idf_keywords', {})
                page_topics[r['url']] = topics if topics is not None else {}

        assigned = set()
        for url1, topics1 in page_topics.items():
            if url1 in assigned:
                continue
            cluster_name = list(topics1.keys())[0] if topics1 else 'other'
            cluster = {'name': cluster_name, 'hub': url1, 'satellites': []}
            assigned.add(url1)
            for url2, topics2 in page_topics.items():
                if url2 in assigned or url1 == url2:
                    continue
                keys1, keys2 = set(topics1.keys()), set(topics2.keys())
                if keys1 and keys2:
                    sim = len(keys1 & keys2) / len(keys1 | keys2)
                    if sim > 0.3:
                        cluster['satellites'].append(url2)
                        assigned.add(url2)
            if cluster['satellites'] or len(topics1) > 0:
                self.topic_clusters[cluster_name] = cluster
        for result in self.results:
            for cluster_name, cluster_data in self.topic_clusters.items():
                if result['url'] == cluster_data['hub']:
                    result['is_topic_hub'] = True
                    result['topic_cluster'] = cluster_name

    def calculate_linking_quality_score(self):
        for result in self.results:
            score = 100
            issues = []
            quality = result.get('anchor_text_quality_score', 0)
            if quality < 50:
                score -= 30
                issues.append('❌ Generic')
            elif quality < 70:
                score -= 15
            total = result.get('total_links', 0)
            if total == 0:
                score -= 20
                issues.append('❌ No links')
            elif total > 50:
                score -= 10
            elif total < 5:
                score -= 15
                issues.append('⚠️ Few')
            result['linking_quality_score'] = max(0, score)
            result['linking_issues'] = issues

    # ==================== КРАУЛИНГ ====================

    def crawl(self):
        """Краулинг с анимацией молнии"""
        self.lightning.start_animation('lightning')
        
        self.load_robots_txt()
        self.load_sitemap()
        
        self.lightning.stop_animation()
        print("\n🕷️ Начинаю краулинг сайта...\n")
        self.lightning.start_animation('bars')
        
        page_count = 0
        while self.to_visit and len(self.visited) < self.max_pages:
            url, depth = self.to_visit.pop(0)
            
            if url in self.visited or depth > self.max_depth:
                continue
            
            if not self.is_valid_url_to_crawl(url):
                continue
            
            if not self.is_url_allowed(url):
                continue
            
            self.visited.add(url)
            page_count += 1
            
            status = f"Анализирую: {url[:50]}..."
            self.lightning.update_status(status)
            self.lightning.print_progress_bar(page_count, min(len(self.visited), self.max_pages), status)
            
            try:
                response = self.session.get(url, timeout=10, allow_redirects=True)
                response.encoding = 'utf-8'
                
                if 'text/html' not in response.headers.get('content-type', '').lower():
                    continue
                
                soup = BeautifulSoup(response.content, 'html.parser')
                analysis = self.analyze_page(soup, url, response)
                self.results.append(analysis)
                
                external = self.extract_external_links(soup, url)
                self.external_links.extend(external)
                
                broken = self.detect_broken_internal_links(soup, url)
                if broken:
                    self.broken_links.extend(broken)
                
                if depth < self.max_depth:
                    for link in soup.find_all('a', href=True):
                        try:
                            next_url = urljoin(url, link['href'])
                            if (self.is_valid_url_to_crawl(next_url) and
                                next_url not in self.visited and
                                len(self.to_visit) < self.max_pages * 2):
                                self.to_visit.append((next_url, depth + 1))
                        except:
                            pass
                
                time.sleep(0.2)
            except:
                pass
        
        self.lightning.stop_animation()
        print(f"\n✅ Краулинг завершён: {len(self.results)} страниц проанализировано\n")
        
        print("🔥 Вычисляю рейтинги и кластеры...")
        self.lightning.start_animation('dots')
        
        self._calculate_internal_pagerank()
        self._calculate_tf_idf()
        self.calculate_site_health_scores()
        self.build_semantic_linking_map()
        self.analyze_anchor_text_quality()
        self.cluster_by_topics()
        self.calculate_linking_quality_score()
        
        self.lightning.stop_animation()
        print("✅ Анализ завершён!\n")

    def analyze_page(self, soup, url, response):
        """ПОЛНЫЙ анализ страницы"""
        text = soup.get_text(separator=' ', strip=True)
        self.all_urls_data[url] = text
        
        for link in soup.find_all('a', href=True):
            try:
                next_url = urljoin(url, link['href']).split('#')[0]
                if urlparse(next_url).netloc == self.domain and self.is_valid_url_to_crawl(next_url):
                    self.internal_links_graph[url].append(next_url)
            except:
                pass
        
        h_hierarchy, h_errors, h_details = self.analyze_h_hierarchy_detailed(soup)
        all_issues = self.collect_all_issues(soup, text, h_errors)
        
        https_ok = self.check_https(response)
        mobile_friendly = self.check_mobile_friendly(soup)
        structured_total, structured_data = self.check_structured_data(soup)
        hreflang = self.check_hreflang(soup)
        breadcrumbs = self.check_breadcrumbs(soup)
        meta_robots = self.check_meta_robots(soup)
        last_modified = self.check_last_modified(response)
        compression = self.check_compression(response)
        cache_control = self.check_cache_headers(response)
        images_opt = self.analyze_images_optimization(soup)
        freshness = self.calculate_content_freshness(response)
        follow_count, nofollow_count = self.count_follow_nofollow(soup)
        
        return {
            'url': url,
            'status': response.status_code,
            'title': soup.title.string if soup.title else '',
            'title_len': len(soup.title.string) if soup.title else 0,
            'description': soup.find('meta', {'name': 'description'})['content'] if soup.find('meta', {'name': 'description'}) else '',
            'desc_len': len(soup.find('meta', {'name': 'description'})['content']) if soup.find('meta', {'name': 'description'}) else 0,
            'h1_count': len(soup.find_all('h1')),
            'h1_text': soup.find('h1').string if soup.find('h1') else '',
            'h_hierarchy': h_hierarchy,
            'h_errors': h_errors,
            'h_details': h_details,
            'heading_distribution': self.analyze_heading_distribution(soup),
            'words_count': len(text.split()),
            'unique_percent': self.calculate_unique_percent(text),
            'boilerplate_percent': self.calculate_boilerplate(text),
            'readability_score': self.analyze_readability(text),
            'avg_sentence_length': self.get_avg_sentence_length(text),
            'avg_word_length': self.get_avg_word_length(text),
            'complex_words_percent': self.count_complex_words(text),
            'content_density': self.calculate_content_density(soup, text),
            'keyword_stuffing_score': self.detect_keyword_stuffing(text),
            'toxicity_score': self.calculate_toxicity_score(text),
            'ai_markers': self.detect_ai_markers(text),
            'filler_phrases': self.count_filler_phrases(text),
            'images_count': len(soup.find_all('img')),
            'images_no_alt': len([img for img in soup.find_all('img') if not img.get('alt')]),
            'int_links': len(soup.find_all('a', href=True)),
            'semantic_tags_count': self.count_semantic_tags(soup),
            'canonical': 1 if soup.find('link', {'rel': 'canonical'}) else 0,
            'schema': 1 if soup.find('script', {'type': 'application/ld+json'}) else 0,
            'og_tags': self.count_og_tags(soup),
            'js_dependence': self.check_js_dependence(soup),
            'dom_nodes': self.count_dom_nodes(soup),
            'has_main_tag': 1 if soup.find('main') else 0,
            'html_quality_score': self.calculate_html_quality_score(soup),
            'deprecated_tags': self.detect_deprecated_tags(soup),
            'hidden_content': self.detect_hidden_content(soup),
            'cloaking_detected': self.detect_cloaking(soup, text),
            'has_contact_info': self.detect_contact_info(soup, text),
            'has_legal_docs': self.detect_legal_docs(soup),
            'has_author_info': self.detect_author_info(soup),
            'has_reviews': self.detect_reviews(soup),
            'trust_badges': self.detect_trust_badges(soup),
            'trust_score': self.calculate_trust_score(soup, text),
            'eeat_score': self.calculate_eeat_score(soup, text),
            'eeat_components': self.analyze_eeat_components(soup, text),
            'cta_count': self.count_ctas(soup),
            'cta_text_quality': self.evaluate_cta_text(soup),
            'lists_count': len(soup.find_all(['ul', 'ol'])),
            'tables_count': len(soup.find_all('table')),
            'faq_count': self.count_faq(soup),
            'site_health_score': 0,
            'top_keywords': self.extract_top_keywords(text),
            'keyword_density_profile': self.get_keyword_density_profile(text),
            'tf_idf_keywords': {},
            'page_authority': 0,
            'incoming_links_count': 0,
            'outgoing_links_internal': len([l for l in soup.find_all('a', href=True) if self.is_valid_url_to_crawl(urljoin(url, l['href']))]),
            'is_orphan': False,
            'semantic_links': [],
            'is_topic_hub': False,
            'topic_cluster': None,
            'anchor_text_quality_score': 0,
            'total_links': 0,
            'linking_quality_score': 0,
            'all_issues': all_issues,
            'https_ok': https_ok,
            'mobile_friendly': mobile_friendly,
            'structured_data': structured_total,
            'structured_data_detail': structured_data,
            'hreflang': hreflang,
            'breadcrumbs': breadcrumbs,
            'meta_robots': meta_robots,
            'last_modified': last_modified,
            'compression': compression,
            'cache_control': cache_control,
            'images_optimization': images_opt,
            'content_freshness_days': freshness,
            'follow_links': follow_count,
            'nofollow_links': nofollow_count,
        }

    # ==================== ФОРМАТИРОВАНИЕ EXCEL ====================

    def get_status_color(self, score, threshold_low=60, threshold_high=80):
        if score >= threshold_high:
            return (self.colors['green_light'], self.colors['green_dark'], '✅')
        elif score >= threshold_low:
            return (self.colors['yellow_light'], self.colors['yellow_dark'], '⚠️')
        else:
            return (self.colors['red_light'], self.colors['red_dark'], '❌')

    def get_status_color_bool(self, value):
        if value:
            return (self.colors['green_light'], self.colors['green_dark'], '✅')
        else:
            return (self.colors['red_light'], self.colors['red_dark'], '❌')

    def apply_cell_color(self, cell, bg_color, font_color, value, bold=True):
        cell.value = value
        cell.fill = PatternFill(start_color=bg_color, end_color=bg_color, fill_type='solid')
        cell.font = Font(color=font_color, bold=bold, size=10)
        cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    def apply_header_style(self, cell, value):
        cell.value = value
        cell.fill = PatternFill(start_color=self.colors['header_dark'], end_color=self.colors['header_dark'], fill_type='solid')
        cell.font = Font(color=self.colors['header_text'], bold=True, size=11)
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    def auto_width_columns(self, ws, max_width=45):
        for column in ws.columns:
            max_length = 0
            col_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[col_letter].width = min(max_length + 2, max_width)

    # ==================== ГЕНЕРАЦИЯ EXCEL (16 ВКЛАДОК) ====================

    def generate_excel_report(self):
        """Генерирует ПОЛНЫЙ Excel отчёт с 16+ вкладками"""
        filename = f"{self.domain.replace('.', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        wb = Workbook()
        
        if 'Sheet' in wb.sheetnames:
            wb.remove(wb['Sheet'])
        
        # ВКЛАДКА 1: ОСНОВНОЙ ОТЧЁТ
        ws = wb.create_sheet('1. Основной отчёт', 0)
        headers = ['URL', 'Title', 'H1', 'Токсичность', 'Иерархия', 'Статус', 'Проблемы']
        for col, header in enumerate(headers, 1):
            self.apply_header_style(ws.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws.cell(row=row, column=1).value = result['url']
            ws.cell(row=row, column=2).value = result['title'] if result['title'] else '-'
            ws.cell(row=row, column=3).value = result['h1_text'] if result['h1_text'] else '-'
            
            tox = result['toxicity_score']
            bg, fg, icon = self.get_status_color(100 - tox, 30, 70)
            self.apply_cell_color(ws.cell(row=row, column=4), bg, fg, f"{icon} {tox:.0f}")
            
            hier = result['h_hierarchy']
            bg, fg = (self.colors['green_light'], self.colors['green_dark']) if hier == 'Good' else (self.colors['red_light'], self.colors['red_dark'])
            self.apply_cell_color(ws.cell(row=row, column=5), bg, fg, hier)
            
            health = result['site_health_score']
            bg, fg, icon = self.get_status_color(health)
            self.apply_cell_color(ws.cell(row=row, column=6), bg, fg, f"{icon} {health:.0f}")
            
            issues = result.get('all_issues', [])[:2]
            ws.cell(row=row, column=7).value = '\n'.join(issues) if issues else '✅'
        
        self.auto_width_columns(ws)
        
        # ВКЛАДКА 2: ОШИБКИ ИЕРАРХИИ
        ws_h = wb.create_sheet('2. Ошибки иерархии', 1)
        h_errors = [r for r in self.results if r['h_hierarchy'] != 'Good']
        
        headers_h = ['URL', 'Статус', 'Проблема', 'Всего заголовков', 'H1 Count', 'Решение']
        for col, header in enumerate(headers_h, 1):
            self.apply_header_style(ws_h.cell(row=1, column=col), header)
        
        row_h = 2
        for result in h_errors:
            ws_h.cell(row=row_h, column=1).value = result['url']
            
            status = result['h_hierarchy']
            if status == 'Bad (wrong start)':
                bg, fg = self.colors['red_light'], self.colors['red_dark']
            elif status == 'Bad (hierarchy broken)':
                bg, fg = self.colors['red_light'], self.colors['red_dark']
            else:
                bg, fg = self.colors['yellow_light'], self.colors['yellow_dark']
            
            self.apply_cell_color(ws_h.cell(row=row_h, column=2), bg, fg, status)
            
            issue_text = '\n'.join(result['h_errors'])
            ws_h.cell(row=row_h, column=3).value = issue_text
            
            details = result.get('h_details', {})
            ws_h.cell(row=row_h, column=4).value = details.get('total_headers', 0)
            ws_h.cell(row=row_h, column=5).value = details.get('h1_count', 0)
            
            if 'wrong start' in status:
                solution = '➡️ Добавьте <h1> в начало основного контента'
            elif 'broken' in status:
                solution = '➡️ Исправьте пропуски уровней (H1→H2→H3)'
            elif 'multiple' in status:
                solution = '➡️ Оставьте только 1 H1, остальные замените на H2'
            else:
                solution = '✅ OK'
            
            ws_h.cell(row=row_h, column=6).value = solution
            row_h += 1
        
        self.auto_width_columns(ws_h, 80)
        
        # ВКЛАДКА 3: ON-PAGE SEO
        ws_onpage = wb.create_sheet('3. On-Page SEO', 2)
        headers_onpage = ['URL', 'Title Len', 'Meta Len', 'H1', 'Canonical', 'Mobile', 'Schema', 'Breadcrumbs']
        for col, header in enumerate(headers_onpage, 1):
            self.apply_header_style(ws_onpage.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_onpage.cell(row=row, column=1).value = result['url']
            ws_onpage.cell(row=row, column=2).value = f"{result['title_len']} ch"
            ws_onpage.cell(row=row, column=3).value = f"{result['desc_len']} ch"
            ws_onpage.cell(row=row, column=4).value = result['h1_count']
            ws_onpage.cell(row=row, column=5).value = '✅' if result['canonical'] else '❌'
            ws_onpage.cell(row=row, column=6).value = '✅' if result['mobile_friendly'] else '❌'
            ws_onpage.cell(row=row, column=7).value = result['structured_data']
            ws_onpage.cell(row=row, column=8).value = '✅' if result['breadcrumbs'] else '❌'
        
        self.auto_width_columns(ws_onpage)
        
        # ВКЛАДКА 4: CONTENT
        ws_content = wb.create_sheet('4. Content', 3)
        headers_content = ['URL', 'Words', 'Unique %', 'Readability', 'Toxicity', 'AI Markers', 'Filler']
        for col, header in enumerate(headers_content, 1):
            self.apply_header_style(ws_content.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_content.cell(row=row, column=1).value = result['url']
            ws_content.cell(row=row, column=2).value = result['words_count']
            ws_content.cell(row=row, column=3).value = f"{result['unique_percent']:.1f}%"
            
            read = result['readability_score']
            if read:
                bg, fg, icon = self.get_status_color(read)
                self.apply_cell_color(ws_content.cell(row=row, column=4), bg, fg, f"{icon} {read:.0f}")
            
            tox = result['toxicity_score']
            bg, fg, icon = self.get_status_color(100 - tox, 30, 70)
            self.apply_cell_color(ws_content.cell(row=row, column=5), bg, fg, f"{icon} {tox:.0f}")
            
            ws_content.cell(row=row, column=6).value = result['ai_markers']
            ws_content.cell(row=row, column=7).value = result['filler_phrases']
        
        self.auto_width_columns(ws_content)
        
        # ВКЛАДКА 5: TECHNICAL
        ws_tech = wb.create_sheet('5. Technical', 4)
        headers_tech = ['URL', 'DOM', 'HTML Score', 'HTTPS', 'Compression', 'Cache', 'Deprecated']
        for col, header in enumerate(headers_tech, 1):
            self.apply_header_style(ws_tech.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_tech.cell(row=row, column=1).value = result['url']
            ws_tech.cell(row=row, column=2).value = result['dom_nodes']
            
            html = result['html_quality_score']
            bg, fg, icon = self.get_status_color(html)
            self.apply_cell_color(ws_tech.cell(row=row, column=3), bg, fg, f"{icon} {html}")
            
            https = '✅' if result['https_ok'] else '❌'
            bg, fg, _ = self.get_status_color_bool(result['https_ok'])
            self.apply_cell_color(ws_tech.cell(row=row, column=4), bg, fg, https)
            
            ws_tech.cell(row=row, column=5).value = 'Yes' if result['compression'] else 'No'
            ws_tech.cell(row=row, column=6).value = 'Set' if result['cache_control'] != 'not set' else 'No'
            ws_tech.cell(row=row, column=7).value = result['deprecated_tags']
        
        self.auto_width_columns(ws_tech)
        
        # ВКЛАДКА 6: E-E-A-T
        ws_eeat = wb.create_sheet('6. E-E-A-T', 5)
        headers_eeat = ['URL', 'Score', 'Expertise', 'Authority', 'Trust', 'Experience']
        for col, header in enumerate(headers_eeat, 1):
            self.apply_header_style(ws_eeat.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_eeat.cell(row=row, column=1).value = result['url']
            eeat = result['eeat_score']
            bg, fg, icon = self.get_status_color(eeat)
            self.apply_cell_color(ws_eeat.cell(row=row, column=2), bg, fg, f"{icon} {eeat:.0f}")
            
            comp = result['eeat_components']
            ws_eeat.cell(row=row, column=3).value = comp.get('expertise', 0)
            ws_eeat.cell(row=row, column=4).value = comp.get('authoritativeness', 0)
            ws_eeat.cell(row=row, column=5).value = comp.get('trustworthiness', 0)
            ws_eeat.cell(row=row, column=6).value = comp.get('experience', 0)
        
        self.auto_width_columns(ws_eeat)
        
        # ВКЛАДКА 7: TRUST
        ws_trust = wb.create_sheet('7. Trust', 6)
        headers_trust = ['URL', 'Trust Score', 'Contact', 'Legal', 'Reviews', 'Badges']
        for col, header in enumerate(headers_trust, 1):
            self.apply_header_style(ws_trust.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_trust.cell(row=row, column=1).value = result['url']
            trust = result['trust_score']
            bg, fg, icon = self.get_status_color(trust)
            self.apply_cell_color(ws_trust.cell(row=row, column=2), bg, fg, f"{icon} {trust:.0f}")
            ws_trust.cell(row=row, column=3).value = '✅' if result['has_contact_info'] else '❌'
            ws_trust.cell(row=row, column=4).value = '✅' if result['has_legal_docs'] else '❌'
            ws_trust.cell(row=row, column=5).value = '✅' if result['has_reviews'] else '❌'
            ws_trust.cell(row=row, column=6).value = result['trust_badges']
        
        self.auto_width_columns(ws_trust)
        
        # ВКЛАДКА 8: HEALTH
        ws_health = wb.create_sheet('8. Health', 7)
        headers_health = ['URL', 'Health Score', 'Words', 'Unique %', 'Readability']
        for col, header in enumerate(headers_health, 1):
            self.apply_header_style(ws_health.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_health.cell(row=row, column=1).value = result['url']
            health = result['site_health_score']
            bg, fg, icon = self.get_status_color(health)
            self.apply_cell_color(ws_health.cell(row=row, column=2), bg, fg, f"{icon} {health:.0f}")
            ws_health.cell(row=row, column=3).value = result['words_count']
            ws_health.cell(row=row, column=4).value = f"{result['unique_percent']:.1f}%"
            
            read = result['readability_score']
            if read:
                ws_health.cell(row=row, column=5).value = f"{read:.0f}"
        
        self.auto_width_columns(ws_health)
        
        # ВКЛАДКА 9: INTERNAL LINKS
        ws_links = wb.create_sheet('9. Internal Links', 8)
        headers_links = ['URL', 'Authority', 'Incoming', 'Outgoing', 'Is Orphan']
        for col, header in enumerate(headers_links, 1):
            self.apply_header_style(ws_links.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_links.cell(row=row, column=1).value = result['url']
            ws_links.cell(row=row, column=2).value = result['page_authority']
            ws_links.cell(row=row, column=3).value = result['incoming_links_count']
            ws_links.cell(row=row, column=4).value = result['outgoing_links_internal']
            ws_links.cell(row=row, column=5).value = '❌ ORPHAN' if result['is_orphan'] else '✅'
        
        self.auto_width_columns(ws_links)
        
        # ВКЛАДКА 10: IMAGES
        ws_img = wb.create_sheet('10. Images', 9)
        headers_img = ['URL', 'Total', 'No Alt', 'No Width', 'No Lazy', 'Issues']
        for col, header in enumerate(headers_img, 1):
            self.apply_header_style(ws_img.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_img.cell(row=row, column=1).value = result['url']
            ws_img.cell(row=row, column=2).value = result['images_optimization']['total']
            
            no_alt = result['images_optimization']['no_alt']
            ws_img.cell(row=row, column=3).value = f"❌ {no_alt}" if no_alt > 0 else '✅'
            
            no_dims = result['images_optimization']['no_width_height']
            ws_img.cell(row=row, column=4).value = f"⚠️ {no_dims}" if no_dims > 0 else '✅'
            
            no_lazy = result['images_optimization']['no_lazy_load']
            ws_img.cell(row=row, column=5).value = f"⚠️ {no_lazy}" if no_lazy > 0 else '✅'
            
            total_issues = no_alt + no_dims + no_lazy
            ws_img.cell(row=row, column=6).value = total_issues
        
        self.auto_width_columns(ws_img)
        
        # ВКЛАДКА 11: EXTERNAL LINKS
        ws_ext = wb.create_sheet('11. External Links', 10)
        headers_ext = ['URL', 'Total External', 'Follow', 'NoFollow', 'Follow %']
        for col, header in enumerate(headers_ext, 1):
            self.apply_header_style(ws_ext.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_ext.cell(row=row, column=1).value = result['url']
            total = result['follow_links'] + result['nofollow_links']
            ws_ext.cell(row=row, column=2).value = total
            ws_ext.cell(row=row, column=3).value = result['follow_links']
            ws_ext.cell(row=row, column=4).value = result['nofollow_links']
            
            follow_pct = (result['follow_links'] / total * 100) if total > 0 else 0
            ws_ext.cell(row=row, column=5).value = f"{follow_pct:.0f}%"
        
        self.auto_width_columns(ws_ext)
        
        # ВКЛАДКА 12: STRUCTURED DATA
        ws_struct = wb.create_sheet('12. Structured Data', 11)
        headers_struct = ['URL', 'Total', 'JSON-LD', 'Microdata', 'RDFa', 'Hreflang', 'Meta Robots']
        for col, header in enumerate(headers_struct, 1):
            self.apply_header_style(ws_struct.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_struct.cell(row=row, column=1).value = result['url']
            ws_struct.cell(row=row, column=2).value = result['structured_data']
            
            detail = result['structured_data_detail']
            ws_struct.cell(row=row, column=3).value = detail.get('json_ld', 0)
            ws_struct.cell(row=row, column=4).value = detail.get('microdata', 0)
            ws_struct.cell(row=row, column=5).value = detail.get('rdfa', 0)
            ws_struct.cell(row=row, column=6).value = result['hreflang']
            ws_struct.cell(row=row, column=7).value = result['meta_robots']
        
        self.auto_width_columns(ws_struct)
        
        # ВКЛАДКА 13: KEYWORDS & TF-IDF
        ws_kw = wb.create_sheet('13. Keywords & TF-IDF', 12)
        headers_kw = ['URL', 'Top Keywords', 'TF-IDF 1', 'TF-IDF 2', 'TF-IDF 3']
        for col, header in enumerate(headers_kw, 1):
            self.apply_header_style(ws_kw.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_kw.cell(row=row, column=1).value = result['url']
            
            top_kw = ', '.join([kw for kw, count in result['top_keywords'][:5]])
            ws_kw.cell(row=row, column=2).value = top_kw
            
            tf_idf = result['tf_idf_keywords']
            tfidf_keys = list(tf_idf.keys())
            ws_kw.cell(row=row, column=3).value = tfidf_keys[0] if len(tfidf_keys) > 0 else '-'
            ws_kw.cell(row=row, column=4).value = tfidf_keys[1] if len(tfidf_keys) > 1 else '-'
            ws_kw.cell(row=row, column=5).value = tfidf_keys[2] if len(tfidf_keys) > 2 else '-'
        
        self.auto_width_columns(ws_kw)
        
        # ВКЛАДКА 14: TOPICS & CLUSTERS
        ws_topics = wb.create_sheet('14. Topics', 13)
        headers_topics = ['URL', 'Is Hub', 'Cluster', 'Incoming Links', 'Semantic Links']
        for col, header in enumerate(headers_topics, 1):
            self.apply_header_style(ws_topics.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_topics.cell(row=row, column=1).value = result['url']
            ws_topics.cell(row=row, column=2).value = '⭐ HUB' if result['is_topic_hub'] else '-'
            ws_topics.cell(row=row, column=3).value = result['topic_cluster'] or '-'
            ws_topics.cell(row=row, column=4).value = result['incoming_links_count']
            
            sem_links = len(result['semantic_links'])
            ws_topics.cell(row=row, column=5).value = f"{sem_links} рекомендаций"
        
        self.auto_width_columns(ws_topics)
        
        # ВКЛАДКА 15: ADVANCED
        ws_adv = wb.create_sheet('15. Advanced', 14)
        headers_adv = ['URL', 'Freshness Days', 'Hidden Content', 'Cloaking', 'CTA Count', 'List/Tables']
        for col, header in enumerate(headers_adv, 1):
            self.apply_header_style(ws_adv.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_adv.cell(row=row, column=1).value = result['url']
            
            fresh = result['content_freshness_days']
            ws_adv.cell(row=row, column=2).value = f"{fresh} дней" if fresh else 'Unknown'
            
            hidden = result['hidden_content']
            ws_adv.cell(row=row, column=3).value = f"❌ {hidden}" if hidden > 0 else '✅'
            
            cloak = '❌' if result['cloaking_detected'] else '✅'
            ws_adv.cell(row=row, column=4).value = cloak
            
            ws_adv.cell(row=row, column=5).value = result['cta_count']
            ws_adv.cell(row=row, column=6).value = f"{result['lists_count']}/{result['tables_count']}"
        
        self.auto_width_columns(ws_adv)
        
        # ВКЛАДКА 16: LINKING QUALITY
        ws_lq = wb.create_sheet('16. Link Quality', 15)
        headers_lq = ['URL', 'Quality Score', 'Total Links', 'Anchor Quality', 'Issues']
        for col, header in enumerate(headers_lq, 1):
            self.apply_header_style(ws_lq.cell(row=1, column=col), header)
        
        for row, result in enumerate(self.results, 2):
            ws_lq.cell(row=row, column=1).value = result['url']
            
            score = result['linking_quality_score']
            bg, fg, icon = self.get_status_color(score)
            self.apply_cell_color(ws_lq.cell(row=row, column=2), bg, fg, f"{icon} {score}")
            
            ws_lq.cell(row=row, column=3).value = result.get('total_links', 0)
            ws_lq.cell(row=row, column=4).value = f"{result['anchor_text_quality_score']:.0f}%"
            
            issues = '\n'.join(result.get('linking_issues', []))
            ws_lq.cell(row=row, column=5).value = issues if issues else '✅'
        
        self.auto_width_columns(ws_lq)
        
        wb.save(filename)
        return filename

    def generate_word_report(self):
        """Генерирует WORD отчёт с 9 разделами"""
        doc = Document()
        
        title = doc.add_paragraph()
        title_run = title.add_run('🚀 SEO Technical Audit Report')
        title_run.font.size = Pt(24)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph(f'Website: {self.base_url}')
        doc.add_paragraph(f'Report Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Pages Analyzed: {len(self.results)}')
        
        critical_issues = [r for r in self.results if r['h_hierarchy'] != 'Good' or 
                          r['unique_percent'] < 50 or r['hidden_content'] > 0]
        
        # РАЗДЕЛ 1
        doc.add_heading('🔴 1. КРИТИЧНЫЕ ОШИБКИ (СРОЧНО ИСПРАВИТЬ)', level=1)
        if critical_issues:
            for result in critical_issues[:10]:
                doc.add_heading(f"URL: {result['url']}", level=2)
                issues_text = []
                if result['h_hierarchy'] != 'Good':
                    issues_text.append(f"❌ ОШИБКА: {result['h_hierarchy']}")
                    if result['h_errors']:
                        issues_text.append(f"   {result['h_errors'][0]}")
                if result['unique_percent'] < 50:
                    issues_text.append(f"❌ Уникальность: {result['unique_percent']:.0f}% (нужно 50%+)")
                if result['hidden_content'] > 0:
                    issues_text.append(f"❌ Скрытый контент: {result['hidden_content']} элементов")
                for issue in issues_text:
                    doc.add_paragraph(issue, style='List Bullet')
        else:
            doc.add_paragraph('✅ Критичных ошибок не найдено!')
        
        # РАЗДЕЛ 2
        doc.add_heading('1️⃣ 2. ИСПРАВЛЕНИЕ ИЕРАРХИИ ЗАГОЛОВКОВ', level=1)
        doc.add_paragraph('Правильная иерархия: H1 → H2 → H3 → H4 → H5 → H6')
        doc.add_paragraph('❌ ОШИБКА #1: Иерархия начинается с H3 вместо H1')
        doc.add_paragraph('Решение: Добавьте <h1> в начало основного контента страницы', style='List Bullet')
        doc.add_paragraph('❌ ОШИБКА #2: Прыжок H2 → H4 (пропущены уровни)')
        doc.add_paragraph('Решение: Используйте H2 → H3 → H4 (последовательно)', style='List Bullet')
        doc.add_paragraph('❌ ОШИБКА #3: Несколько H1 на одной странице')
        doc.add_paragraph('Решение: Оставьте только 1 H1, остальные замените на H2', style='List Bullet')
        
        # РАЗДЕЛ 3
        doc.add_heading('2️⃣ 3. ОПТИМИЗАЦИЯ КОНТЕНТА', level=1)
        doc.add_paragraph(f'Проанализировано слов: {sum(r["words_count"] for r in self.results)} всего')
        doc.add_paragraph(f'Среднее: {int(sum(r["words_count"] for r in self.results) / len(self.results))} слов на страницу')
        doc.add_paragraph('Требования:')
        doc.add_paragraph('✅ Минимум 300 слов на странице', style='List Bullet')
        doc.add_paragraph('✅ Уникальность минимум 50%', style='List Bullet')
        doc.add_paragraph('✅ Читаемость от 60 баллов', style='List Bullet')
        doc.add_paragraph('✅ Без keyword stuffing (максимум 3% плотность)', style='List Bullet')
        
        # РАЗДЕЛ 4
        doc.add_heading('3️⃣ 4. ОПТИМИЗАЦИЯ КАРТИНОК', level=1)
        img_issues = sum(r['images_no_alt'] for r in self.results)
        doc.add_paragraph(f'Всего картинок без ALT: {img_issues}')
        doc.add_paragraph('Решение: Добавьте описательный ALT текст')
        doc.add_paragraph('Правила ALT текста:')
        doc.add_paragraph('✅ Описывайте, что на картинке (не просто "картинка")', style='List Bullet')
        doc.add_paragraph('✅ Используйте ключевые слова (естественно)', style='List Bullet')
        doc.add_paragraph('✅ Максимум 125 символов', style='List Bullet')
        doc.add_paragraph('✅ Не повторяйте текст рядом', style='List Bullet')
        
        # РАЗДЕЛ 5
        doc.add_heading('4️⃣ 5. TECHNICAL SEO', level=1)
        doc.add_paragraph(f'HTTPS: {sum(1 for r in self.results if r["https_ok"])} из {len(self.results)} ({round(sum(1 for r in self.results if r["https_ok"]) / len(self.results) * 100, 1)}%)')
        doc.add_paragraph(f'Mobile Friendly: {sum(1 for r in self.results if r["mobile_friendly"])} из {len(self.results)}')
        doc.add_paragraph(f'Schema Markup: {sum(1 for r in self.results if r["schema"])} из {len(self.results)}')
        doc.add_paragraph(f'Canonical Tags: {sum(1 for r in self.results if r["canonical"])} из {len(self.results)}')
        
        # РАЗДЕЛ 6
        doc.add_heading('5️⃣ 6. ON-PAGE SEO', level=1)
        doc.add_paragraph('Title Tag:')
        doc.add_paragraph('✅ Длина: 30-60 символов', style='List Bullet')
        doc.add_paragraph('✅ Включите основной ключевое слово', style='List Bullet')
        doc.add_paragraph('Meta Description:')
        doc.add_paragraph('✅ Длина: 100-160 символов', style='List Bullet')
        doc.add_paragraph('✅ Включите призыв к действию', style='List Bullet')
        doc.add_paragraph('H1 Tag:')
        doc.add_paragraph('✅ Только 1 на странице', style='List Bullet')
        doc.add_paragraph('✅ Должен быть около начала контента', style='List Bullet')
        
        # РАЗДЕЛ 7
        doc.add_heading('6️⃣ 7. КЛЮЧЕВЫЕ СЛОВА И СЕМАНТИКА', level=1)
        all_keywords = []
        for r in self.results:
            all_keywords.extend([kw[0] for kw in r['top_keywords']])
        if all_keywords:
            top_kw = Counter(all_keywords).most_common(15)
            doc.add_paragraph('TOP 15 ключевых слов на сайте:')
            for kw, count in top_kw:
                doc.add_paragraph(f'{kw}: {count} упоминаний', style='List Bullet')
        
        # РАЗДЕЛ 8
        doc.add_heading('7️⃣ 8. ДОВЕРИЕ И E-E-A-T', level=1)
        doc.add_paragraph('Компоненты E-E-A-T:')
        doc.add_paragraph('Expertise (Экспертиза):', style='List Bullet')
        doc.add_paragraph('✅ Информация об авторе', style='List Bullet')
        doc.add_paragraph('✅ Источники и ссылки на авторитеты', style='List Bullet')
        doc.add_paragraph('Authoritativeness (Авторитетность):', style='List Bullet')
        doc.add_paragraph('✅ Внешние ссылки на авторитеты', style='List Bullet')
        doc.add_paragraph('✅ Упоминания в средствах массовой информации', style='List Bullet')
        doc.add_paragraph('Trustworthiness (Надежность):', style='List Bullet')
        doc.add_paragraph('✅ Контакты и информация о компании', style='List Bullet')
        doc.add_paragraph('✅ Политика конфиденциальности', style='List Bullet')
        doc.add_paragraph('✅ Отзывы клиентов', style='List Bullet')
        
        # РАЗДЕЛ 9
        doc.add_heading('8️⃣ 9. РЕКОМЕНДАЦИИ И ДЕЙСТВИЯ', level=1)
        doc.add_paragraph('Срочные действия (неделя 1):', style='List Number')
        doc.add_paragraph('1. Исправьте иерархию H1-H6 на всех странах', style='List Bullet')
        doc.add_paragraph('2. Добавьте ALT текст ко всем картинкам', style='List Bullet')
        doc.add_paragraph('3. Проверьте HTTPS на всех страницах', style='List Bullet')
        doc.add_paragraph()
        doc.add_paragraph('Средне-срочные действия (неделя 2-3):', style='List Number')
        doc.add_paragraph('1. Увеличьте уникальность контента до 50%+', style='List Bullet')
        doc.add_paragraph('2. Добавьте Schema Markup (JSON-LD)', style='List Bullet')
        doc.add_paragraph('3. Оптимизируйте Title и Meta Description', style='List Bullet')
        doc.add_paragraph()
        doc.add_paragraph('Долго-срочные действия (месяц 1-3):', style='List Number')
        doc.add_paragraph('1. Улучшите E-E-A-T (Expertise, Authority, Trust)', style='List Bullet')
        doc.add_paragraph('2. Добавьте отзывы и рейтинги', style='List Bullet')
        doc.add_paragraph('3. Создайте Семантическую сетку ссылок', style='List Bullet')
        
        filename = f"{self.domain.replace('.', '_')}_DETAILED_REPORT_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        return filename

    def generate_reports(self):
        """Генерирует оба отчёта"""
        print("📊 Генерирую ПОЛНЫЙ Excel отчёт (16 вкладок)...")
        excel_file = self.generate_excel_report()
        
        print("📄 Генерирую ПОЛНЫЙ Word отчёт (9 разделов)...")
        word_file = self.generate_word_report()
        
        return excel_file, word_file


def main():
    lightning = LightningAnimation()
    
    # Печать приветствия
    lightning.print_title()
    
    if len(sys.argv) < 2:
        print("🚀 SEO Audit Parser v9.0 COMPLETE RESTORED (85+ КБ)")
        print("Использование: python seo_audit_v9_0.py <URL> [max_pages] [max_depth]\n")
        sys.exit(1)
    
    url = sys.argv[1]
    max_pages = int(sys.argv[2]) if len(sys.argv) > 2 else 50
    max_depth = int(sys.argv[3]) if len(sys.argv) > 3 else 3
    
    print(f"📊 URL: {url}")
    print(f"📄 Max Pages: {max_pages}")
    print(f"📐 Max Depth: {max_depth}")
    lightning.print_divider()
    print()
    
    audit = SEOAuditParser(url, max_pages, max_depth)
    audit.crawl()
    
    print("📊 Генерирую отчёты...\n")
    excel_file, word_file = audit.generate_reports()
    
    lightning.print_divider()
    print("\n✅ AUDIT COMPLETE!")
    print(f"📊 Excel (16 вкладок): {excel_file}")
    print(f"📄 Word (9 разделов): {word_file}")
    print(f"\nEXCEL ВКЛАДКИ:")
    for i in range(1, 17):
        names = [
            'Основной отчёт', 'Ошибки иерархии', 'On-Page SEO', 'Content',
            'Technical', 'E-E-A-T', 'Trust', 'Health', 'Internal Links',
            'Images', 'External Links', 'Structured Data', 'Keywords & TF-IDF',
            'Topics', 'Advanced', 'Link Quality'
        ]
        print(f"{i}. {names[i-1]}")
    print(f"\nWORD РАЗДЕЛЫ:")
    print("1. Критичные ошибки")
    print("2. Иерархия заголовков")
    print("3. Контент")
    print("4. Картинки")
    print("5. Technical SEO")
    print("6. On-Page SEO")
    print("7. Ключевые слова")
    print("8. Доверие и E-E-A-T")
    print("9. Рекомендации")
    lightning.print_divider()
    print()


if __name__ == '__main__':
    main()
